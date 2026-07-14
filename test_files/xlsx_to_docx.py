#!/usr/bin/env python3
"""Render the Universal Data Suite test report .xlsx into a clean, official
looking .docx (one section/page per worksheet).

It does NOT mirror the spreadsheet grid cell-by-cell. Instead it reads the
*content and meaning* of each sheet and lays it out as a document:

  * a key/value sheet (no 'S.No' header, e.g. Summary) becomes a title, an
    intro line, a borderless label/value block and a signature area;
  * a tabular sheet (has an 'S.No' header row, e.g. Requirements / Test Cases)
    becomes a heading + a neat table: coloured header, subtle row striping,
    thin borders only around the real data, green/red status cells.

Colours (navy headers, green PASS) and emphasis are kept; empty spreadsheet
rows and full gridlines are dropped.

Usage: python xlsx_to_docx.py <input.xlsx> <output.docx>
"""
import sys
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INP = sys.argv[1] if len(sys.argv) > 1 else 'test_files/output/Universal_Data_Suite_Test_Report.xlsx'
OUT = sys.argv[2] if len(sys.argv) > 2 else 'test_files/output/Universal_Data_Suite_Test_Report.docx'

NAVY = '1F3864'
WHITE = 'FFFFFF'
ZEBRA = 'F2F7FF'
GREENF = 'C6EFCE'
GREENT = '006100'
REDF = 'FFC7CE'
REDT = '9C0006'
GREY = '595959'
LINE = 'BFBFBF'

DEFAULT_COLW = 8.43


# ----------------------------------------------------------------- helpers
def rgb_of(color):
    try:
        if color is not None and color.type == 'rgb' and color.rgb:
            return str(color.rgb)[-6:]
    except Exception:
        pass
    return None


def _set(el, tag, **attrs):
    child = OxmlElement(tag)
    for k, v in attrs.items():
        child.set(qn(k.replace('_', ':')), v)
    el.append(child)
    return child


def shade(cell, hexrgb):
    if not hexrgb:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    _set(tcPr, 'w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': hexrgb})


def no_borders(table):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        _set(b, 'w:' + side, **{'w:val': 'nil'})
    tblPr.append(b)


def grid_borders(table, color=LINE, sz='4'):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        _set(b, 'w:' + side, **{'w:val': 'single', 'w:sz': sz, 'w:space': '0', 'w:color': color})
    tblPr.append(b)


def cell_margins(table, top=11, bottom=11, left=70, right=70):
    tblPr = table._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        _set(mar, 'w:' + side, **{'w:w': str(val), 'w:type': 'dxa'})
    tblPr.append(mar)


def fixed_layout(table):
    _set(table._tbl.tblPr, 'w:tblLayout', **{'w:type': 'fixed'})


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    _set(trPr, 'w:tblHeader', **{'w:val': 'true'})


def style_run(run, size, color=None, bold=False, italic=False, font='Calibri'):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def fill_cell(cell, text, size, color=None, bold=False, align='left', valign='center', shade_rgb=None):
    cell.vertical_alignment = {'center': WD_ALIGN_VERTICAL.CENTER,
                               'top': WD_ALIGN_VERTICAL.TOP}.get(valign, WD_ALIGN_VERTICAL.CENTER)
    p = cell.paragraphs[0]
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    style_run(p.add_run('' if text is None else str(text)), size, color, bold)
    if shade_rgb:
        shade(cell, shade_rgb)


def section_setup(section, landscape):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Inches(11), Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.4)
    section.left_margin = section.right_margin = Inches(0.6)


def _field(paragraph, code):
    run = paragraph.add_run()
    style_run(run, 8, GREY)
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = code
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def add_footer_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run('Page '), 8, GREY)
    _field(p, 'PAGE')
    style_run(p.add_run(' of '), 8, GREY)
    _field(p, 'NUMPAGES')


def title_block(doc, title, subtitle=None, big=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if big else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run(title), 17 if big else 13, NAVY, bold=True, font='Calibri Light')
    # navy underline rule
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    _set(pbdr, 'w:bottom', **{'w:val': 'single', 'w:sz': '12', 'w:space': '4', 'w:color': NAVY})
    pPr.append(pbdr)
    if subtitle:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after = Pt(10)
        style_run(s.add_run(subtitle), 9.5, GREY, italic=True)


def used_columns(ws, header_row):
    cols = []
    for c in range(1, ws.max_column + 1):
        if ws.cell(header_row, c).value not in (None, ''):
            cols.append(c)
    return cols


def find_header_row(ws):
    """A tabular sheet has a row containing an 'S.No' header cell; return it."""
    for r in range(1, min(ws.max_row, 12) + 1):
        for c in range(1, ws.max_column + 1):
            v = ws.cell(r, c).value
            if v is not None and str(v).strip().lower() == 's.no':
                return (r, c)
    return None


# ----------------------------------------------------------------- builders
def build_table_sheet(doc, ws, header_row, landscape):
    # sheet title = the merged banner above the header (row 1), else the name
    banner = None
    for c in range(1, ws.max_column + 1):
        v = ws.cell(1, c).value
        if v and str(v).strip():
            banner = str(v).strip(); break
    title_block(doc, banner or ws.title, big=False)

    cols = used_columns(ws, header_row)
    ncols = len(cols)
    headers = [str(ws.cell(header_row, c).value).strip() for c in cols]
    status_idx = next((i for i, h in enumerate(headers) if h.lower() == 'status'), None)
    center_idx = {i for i, h in enumerate(headers)
                  if h.lower() in ('s.no', 'test id', 'status')}

    # data rows = rows below header that carry real content
    data_rows = []
    for r in range(header_row + 1, ws.max_row + 1):
        vals = [ws.cell(r, c).value for c in cols]
        if any(v not in (None, '') for v in vals):
            data_rows.append([('' if v is None else str(v)) for v in vals])

    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    fixed_layout(table)
    grid_borders(table)
    cell_margins(table)

    # column widths proportional to the spreadsheet, scaled to the page
    usable = (11 - 1.4) if landscape else (8.5 - 1.4)
    widths = []
    for c in cols:
        dim = ws.column_dimensions.get(get_column_letter(c))
        widths.append(dim.width if (dim and dim.width) else DEFAULT_COLW)
    tot = sum(widths) or 1.0
    col_in = [usable * w / tot for w in widths]

    body_sz = 7.5 if landscape else 9.5
    # header
    hrow = table.rows[0]
    repeat_header(hrow)
    for i, c in enumerate(cols):
        cell = hrow.cells[i]
        cell.width = Inches(col_in[i])
        fill_cell(cell, headers[i], body_sz + 0.5, WHITE, bold=True, align='center', shade_rgb=NAVY)
    # data
    for ri, row in enumerate(data_rows):
        zebra = ZEBRA if ri % 2 == 1 else None
        for i, text in enumerate(row):
            cell = table.rows[ri + 1].cells[i]
            cell.width = Inches(col_in[i])
            align = 'center' if i in center_idx else 'left'
            if status_idx is not None and i == status_idx and text:
                passed = text.strip().lower().startswith('pass')
                fill_cell(cell, text, body_sz, GREENT if passed else REDT, bold=True,
                          align='center', shade_rgb=GREENF if passed else REDF)
            else:
                fill_cell(cell, text, body_sz, align=align, valign='top', shade_rgb=zebra)

    return landscape


def build_summary_sheet(doc, ws):
    # classify the non-empty content
    title, subtitle = None, None
    title_sz = 0
    meta = []          # (label, value, value_fill)
    sig_header_row = None
    sig_cols = {}

    for r in range(1, ws.max_row + 1):
        vals = [ws.cell(r, c).value for c in range(1, ws.max_column + 1)]
        nonempty = [(c + 1, ws.cell(r, c + 1)) for c in range(ws.max_column)
                    if ws.cell(r, c + 1).value not in (None, '')]
        if not nonempty:
            continue
        texts = [str(cell.value).strip() for _, cell in nonempty]
        low = [t.lower() for t in texts]
        if 'developer' in low or 'tester' in low:
            sig_header_row = r
            sig_cols = {c: str(cell.value).strip() for c, cell in nonempty}
            break  # signature block (and anything below) handled separately
        if len(nonempty) == 1:
            cell = nonempty[0][1]
            sz = (cell.font.size or 11) if cell.font else 11
            if sz >= 14 and not title:
                title, title_sz = texts[0], sz
            elif not subtitle and len(texts[0]) > 30:
                subtitle = texts[0]
        else:
            label = texts[0]
            value = texts[-1]
            vfill = None
            vcell = nonempty[-1][1]
            if vcell.fill is not None and vcell.fill.patternType:
                vfill = rgb_of(vcell.fill.fgColor)
            meta.append((label, value, vfill))

    title_block(doc, title or 'Test Report', subtitle)

    # metadata as a clean borderless two-column block
    if meta:
        t = doc.add_table(rows=len(meta), cols=2)
        t.autofit = False
        fixed_layout(table=t)
        no_borders(t)
        cell_margins(t, top=30, bottom=30, left=40, right=80)
        for i, (label, value, vfill) in enumerate(meta):
            lc, vc = t.rows[i].cells
            lc.width = Inches(2.2); vc.width = Inches(4.8)
            fill_cell(lc, label, 10.5, NAVY, bold=True, align='left')
            is_pass = vfill in (GREENF,) or value.strip().lower() in ('all passed', 'passed')
            fill_cell(vc, value, 10.5, GREENT if is_pass else None, bold=is_pass,
                      align='left', shade_rgb=GREENF if is_pass else None)

    # signature block at the bottom
    if sig_header_row:
        names = {c: [] for c in sig_cols}
        for r in range(sig_header_row + 1, ws.max_row + 1):
            for c in sig_cols:
                v = ws.cell(r, c).value
                if v not in (None, ''):
                    names[c].append(str(v).strip())
        doc.add_paragraph().paragraph_format.space_after = Pt(28)
        cols_sorted = sorted(sig_cols)
        sig = doc.add_table(rows=2, cols=len(cols_sorted))
        sig.autofit = False
        fixed_layout(sig)
        no_borders(sig)
        for j, c in enumerate(cols_sorted):
            top = sig.rows[0].cells[j]
            # signature line
            p = top.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r0 = p.add_run('______________________')
            style_run(r0, 10, GREY)
            r1 = p.add_run('\n' + sig_cols[c])
            style_run(r1, 10.5, NAVY, bold=True)
            bot = sig.rows[1].cells[j]
            bp = bot.paragraphs[0]
            for k, line in enumerate(names[c]):
                run = bp.add_run(('\n' if k else '') + line)
                style_run(run, 10, bold=(k == 0))
    return False  # portrait


def build(inp, out):
    wb = openpyxl.load_workbook(inp)
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    for si, name in enumerate(wb.sheetnames):
        ws = wb[name]
        if si == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section(WD_SECTION.NEW_PAGE)

        hdr = find_header_row(ws)
        if hdr:
            landscape = len(used_columns(ws, hdr[0])) >= 4
            section_setup(section, landscape)
            add_footer_page_number(section)
            build_table_sheet(doc, ws, hdr[0], landscape)
        else:
            section_setup(section, False)
            add_footer_page_number(section)
            build_summary_sheet(doc, ws)

    doc.save(out)
    print('wrote', out)


if __name__ == '__main__':
    build(INP, OUT)
