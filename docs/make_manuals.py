#!/usr/bin/env python3
"""Build both user manuals from their Markdown sources into:
  * an in-app HTML (rendered by the Qt HelpManualDialog, with heading anchors for the TOC + search)
  * a .docx (python-docx, offline)
  * a per-app .qrc embedding the HTML + that app's screenshots under :/manual/

Run offline:  python docs/make_manuals.py
Markdown subset: #/##/### headings, paragraphs, **bold**, `code`, [text](url), pipe tables,
bullet (- ) / numbered (1. ) lists, > blockquotes, ```fenced code```, ![alt](img), --- rules.
"""
import os, re, struct, html as htmllib

HERE = os.path.dirname(os.path.abspath(__file__))
MANUAL = os.path.join(HERE, 'manual')

MANUALS = [
    # (markdown,                  html out,                 docx out,                 title,                                          app qrc,                                          img dir)
    ('simulator_manual.md', 'simulator_manual.html', 'simulator_manual.docx',
     'Universal Data Simulator', os.path.join(HERE, '..', 'app_simulator', 'simulator_manual.qrc'), 'sim'),
    ('parser_manual.md', 'parser_manual.html', 'parser_manual.docx',
     'Universal Wireshark Log Reader', os.path.join(HERE, '..', 'app_parser', 'parser_manual.qrc'), 'parser'),
]

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')
IMG = re.compile(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$')


def slug(text):
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    t = re.sub(r'`([^`]+)`', r'\1', t)
    t = re.sub(r'[^a-zA-Z0-9]+', '-', t).strip('-').lower()
    return t or 'section'


def split_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def is_sep(line):
    cells = split_row(line)
    return len(cells) > 0 and all(re.fullmatch(r':?-{2,}:?', c) for c in cells)


def parse_markdown(path):
    """Parse markdown into a flat list of block tuples."""
    with open(path, 'r', encoding='utf-8') as f:
        lines = [l.rstrip('\n') for l in f]
    blocks, i, n, para = [], 0, len(lines), []

    def flush():
        if para:
            text = ' '.join(s.strip() for s in para).strip()
            if text:
                blocks.append(('p', text))
            para.clear()

    while i < n:
        line, s = lines[i], lines[i].strip()
        if s.startswith('```'):
            flush(); i += 1; code = []
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1; blocks.append(('code', code)); continue
        m = IMG.match(line)
        if m:
            flush(); blocks.append(('img', m.group(1), m.group(2))); i += 1; continue
        mh = re.match(r'^(#{1,6})\s+(.*)$', line)
        if mh:
            flush(); blocks.append(('h', len(mh.group(1)), mh.group(2).strip())); i += 1; continue
        if re.fullmatch(r'(-{3,}|\*{3,}|_{3,})', s):
            flush(); blocks.append(('hr',)); i += 1; continue
        if s.startswith('|') and i + 1 < n and is_sep(lines[i + 1]):
            flush(); header = split_row(line); ncol = len(header); body = []; i += 2
            while i < n and lines[i].strip().startswith('|'):
                cells = split_row(lines[i])
                if len(cells) < ncol: cells += [''] * (ncol - len(cells))
                body.append(cells[:ncol]); i += 1
            blocks.append(('table', header, body)); continue
        if s.startswith('>'):
            flush(); buf = []
            while i < n and lines[i].strip().startswith('>'):
                buf.append(re.sub(r'^\s*>\s?', '', lines[i])); i += 1
            blocks.append(('quote', ' '.join(x.strip() for x in buf).strip())); continue
        if re.match(r'^[-*]\s+', line):
            flush(); items = []
            while i < n and re.match(r'^[-*]\s+', lines[i]):
                content = re.sub(r'^[-*]\s+', '', lines[i]); i += 1
                while i < n and re.match(r'^\s{2,}\S', lines[i]) and not re.match(r'^\s*[-*]\s+', lines[i]) and not re.match(r'^\s*\d+\.\s', lines[i]):
                    content += ' ' + lines[i].strip(); i += 1
                items.append(content)
            blocks.append(('ul', items)); continue
        if re.match(r'^\d+\.\s+', line):
            flush(); items = []
            while i < n and re.match(r'^\d+\.\s+', lines[i]):
                content = re.sub(r'^\d+\.\s+', '', lines[i]); i += 1
                while i < n and re.match(r'^\s{3,}\S', lines[i]) and not re.match(r'^\s*\d+\.\s+', lines[i]) and not re.match(r'^\s*[-*]\s+', lines[i]):
                    content += ' ' + lines[i].strip(); i += 1
                items.append(content)
            blocks.append(('ol', items)); continue
        if s == '':
            flush(); i += 1; continue
        para.append(line); i += 1
    flush()
    return blocks


# ---------------------------------------------------------------- HTML emitter
def inline_html(text):
    out, pos = [], 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            out.append(htmllib.escape(text[pos:m.start()]))
        tok = m.group(0)
        if tok.startswith('**'):
            out.append('<b>' + htmllib.escape(tok[2:-2]) + '</b>')
        elif tok.startswith('`'):
            out.append('<code>' + htmllib.escape(tok[1:-1]) + '</code>')
        else:
            mm = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok)
            out.append('<a href="%s">%s</a>' % (htmllib.escape(mm.group(2)), htmllib.escape(mm.group(1))))
        pos = m.end()
    if pos < len(text):
        out.append(htmllib.escape(text[pos:]))
    return ''.join(out)


# Tuned for Qt's QTextBrowser rich-text engine (a CSS 2.1 subset): spacing is done
# with block margins/padding (line-height is ignored), tables use cell attributes +
# per-cell borders, and the page is a fixed white "document" so it reads the same in
# either app theme. Images are emitted with explicit scaled width/height (Qt ignores
# max-width) so wide screenshots fit the pane instead of forcing a scrollbar.
IMG_MAX_W = 720

CSS = """
body{font-family:'Segoe UI',Arial,sans-serif;font-size:11pt;color:#1F2933;background-color:#FFFFFF;}
h1{font-size:21pt;color:#3730A3;margin-top:2px;margin-bottom:12px;}
h2{font-size:15pt;color:#4338CA;margin-top:24px;margin-bottom:8px;padding-bottom:4px;border-bottom:2px solid #E5E7EB;}
h3{font-size:12.5pt;color:#111827;margin-top:16px;margin-bottom:6px;}
p{margin-top:0px;margin-bottom:10px;}
a{color:#2563EB;text-decoration:none;}
ul,ol{margin-top:0px;margin-bottom:10px;}
li{margin-bottom:4px;}
code{font-family:Consolas,'Courier New',monospace;font-size:10pt;color:#B91C1C;background-color:#F3F4F6;}
pre{font-family:Consolas,'Courier New',monospace;font-size:9.5pt;color:#111827;background-color:#F3F4F6;border:1px solid #E5E7EB;padding:10px;}
blockquote{color:#3730A3;background-color:#EEF2FF;border-left:4px solid #6366F1;padding:8px 14px;margin-left:0px;margin-right:0px;}
table{margin-top:4px;margin-bottom:12px;}
th{color:#1F2933;}
td{color:#1F2933;}
.caption{color:#6B7280;font-style:italic;font-size:9pt;}
hr{color:#E5E7EB;}
"""


def emit_html(blocks, out_path, title):
    parts = ['<!DOCTYPE html><html><head><meta charset="utf-8"><style>%s</style></head><body>' % CSS]
    for b in blocks:
        t = b[0]
        if t == 'h':
            lvl = min(max(b[1], 1), 6); sid = slug(b[2])
            parts.append('<h%d id="%s"><a name="%s"></a>%s</h%d>' % (lvl, sid, sid, inline_html(b[2]), lvl))
        elif t == 'p':
            parts.append('<p>%s</p>' % inline_html(b[1]))
        elif t == 'img':
            # Qt's rich-text engine ignores max-width, so emit an explicit width/height
            # scaled to fit the pane (preserving aspect ratio) — otherwise wide
            # screenshots force a horizontal scrollbar and look "choppy".
            dim = ''
            sz = png_size(os.path.join(MANUAL, b[2]))
            if sz:
                w, h = sz
                if w > IMG_MAX_W:
                    h = max(1, int(round(h * IMG_MAX_W / float(w)))); w = IMG_MAX_W
                dim = ' width="%d" height="%d"' % (w, h)
            parts.append('<p><img src="%s" alt="%s"%s><br><span class="caption">%s</span></p>'
                         % (htmllib.escape(b[2]), htmllib.escape(b[1]), dim, htmllib.escape(b[1])))
        elif t == 'code':
            parts.append('<pre>%s</pre>' % htmllib.escape('\n'.join(b[1])))
        elif t == 'quote':
            parts.append('<blockquote>%s</blockquote>' % inline_html(b[1]))
        elif t == 'ul':
            parts.append('<ul>' + ''.join('<li>%s</li>' % inline_html(x) for x in b[1]) + '</ul>')
        elif t == 'ol':
            parts.append('<ol>' + ''.join('<li>%s</li>' % inline_html(x) for x in b[1]) + '</ol>')
        elif t == 'hr':
            parts.append('<hr>')
        elif t == 'table':
            header, body = b[1], b[2]
            rows = ['<tr>' + ''.join('<th bgcolor="#EEF2FF">%s</th>' % inline_html(c) for c in header) + '</tr>']
            for row in body:
                rows.append('<tr>' + ''.join('<td>%s</td>' % inline_html(c) for c in row) + '</tr>')
            parts.append('<table border="1" bordercolor="#D1D5DB" cellspacing="0" cellpadding="6" width="100%">'
                         + ''.join(rows) + '</table>')
    parts.append('</body></html>')
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(parts))
    return out_path


# ---------------------------------------------------------------- DOCX emitter
def png_size(path):
    try:
        with open(path, 'rb') as f:
            head = f.read(24)
        if head[:8] == b'\x89PNG\r\n\x1a\n':
            return struct.unpack('>II', head[16:24])
    except Exception:
        pass
    return None


def emit_docx(blocks, out_path, img_base, title):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # Roomier page: slightly narrower margins so body text and (smaller) images sit
    # comfortably without dominating the page.
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    IMG_MAX_W_IN = 4.7   # screenshots are scaled to fit this width / height, never full-page
    IMG_MAX_H_IN = 4.2

    def fld(parent_run, instr_text, placeholder):
        # Insert a Word field (e.g. TOC / PAGE) into an existing run element.
        r = parent_run._r
        b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
        i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = instr_text
        s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
        tx = OxmlElement('w:t'); tx.text = placeholder
        e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
        for el in (b, i, s, tx, e):
            r.append(el)

    def add_page_numbers():
        p = doc.sections[0].footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Page ")
        run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        fld(run, "PAGE", "1")

    def add_cover_and_toc():
        # Title (already carries "… — User Manual"), a light subtitle, then a Word
        # Table of Contents field that populates page numbers when the doc opens.
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub = doc.add_paragraph().add_run("Universal Data Suite — user manual")
        sub.italic = True; sub.font.size = Pt(12); sub.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        cap = doc.add_paragraph().add_run("Contents")
        cap.bold = True; cap.font.size = Pt(15); cap.font.color.rgb = RGBColor(0x31, 0x2E, 0x81)
        toc = doc.add_paragraph().add_run()
        fld(toc, 'TOC \\o "1-2" \\h \\z \\u',
            'Update this field (select all, press F9) to build the contents with page numbers.')
        doc.add_page_break()
        # Tell Word to refresh fields (the TOC) on first open.
        st = doc.settings.element
        uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
        st.append(uf)

    add_page_numbers()
    cover_done = [False]

    def add_inline(par, text):
        pos = 0
        for m in INLINE.finditer(text):
            if m.start() > pos:
                par.add_run(text[pos:m.start()])
            tok = m.group(0)
            if tok.startswith('**'):
                par.add_run(tok[2:-2]).bold = True
            elif tok.startswith('`'):
                r = par.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            else:
                mm = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok)
                par.add_run(mm.group(1) + ((' (' + mm.group(2) + ')') if mm.group(2).startswith('http') else ''))
            pos = m.end()
        if pos < len(text):
            par.add_run(text[pos:])

    def shade(cell, hexcolor):
        sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
        cell._tc.get_or_add_tcPr().append(sh)

    for b in blocks:
        t = b[0]
        if t == 'h':
            txt = re.sub(r'\*\*(.+?)\*\*', r'\1', b[2]); txt = re.sub(r'`([^`]+)`', r'\1', txt)
            # The first top-level heading becomes the cover/title + Contents page;
            # every other heading maps to a Word Heading style (1.. for ##, ###).
            if b[1] == 1 and not cover_done[0]:
                add_cover_and_toc()
                cover_done[0] = True
            else:
                doc.add_heading(txt, level=min(max(b[1] - 1, 1), 9))
        elif t == 'p':
            add_inline(doc.add_paragraph(), b[1])
        elif t == 'img':
            fp = os.path.normpath(os.path.join(img_base, b[2]))
            if os.path.exists(fp):
                sz = png_size(fp)
                width = None
                if sz and sz[0] > 0 and sz[1] > 0:
                    w_in, h_in = sz[0] / 96.0, sz[1] / 96.0
                    scale = min(IMG_MAX_W_IN / w_in, IMG_MAX_H_IN / h_in, 1.0)
                    width = Inches(w_in * scale)
                pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    run = pic.add_run()
                    run.add_picture(fp, width=width) if width else run.add_picture(fp)
                except Exception:
                    pic.add_run('[image: %s]' % b[2])
                capp = doc.add_paragraph(); capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = capp.add_run(b[1]); cap.italic = True; cap.font.size = Pt(9)
                cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            else:
                doc.add_paragraph('[missing image: %s]' % b[2])
        elif t == 'code':
            for cl in b[1]:
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
                r = p.add_run(cl if cl else ' '); r.font.name = 'Consolas'; r.font.size = Pt(8.5)
            doc.add_paragraph()
        elif t == 'quote':
            add_inline(doc.add_paragraph(style='Quote'), b[1])
        elif t == 'ul':
            for it in b[1]:
                add_inline(doc.add_paragraph(style='List Bullet'), it)
        elif t == 'ol':
            for k, it in enumerate(b[1], 1):
                p = doc.add_paragraph(); p.add_run('%d. ' % k); add_inline(p, it)
        elif t == 'hr':
            p = doc.add_paragraph(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'BBBBBB')
            pbdr.append(bottom); p._p.get_or_add_pPr().append(pbdr)
        elif t == 'table':
            header, body = b[1], b[2]; ncol = len(header)
            tbl = doc.add_table(rows=1, cols=ncol); tbl.style = 'Table Grid'; tbl.autofit = True
            for c, h in enumerate(header):
                cell = tbl.rows[0].cells[c]; shade(cell, 'DDEBFF'); cell.paragraphs[0].text = ''
                add_inline(cell.paragraphs[0], h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in body:
                cells = tbl.add_row().cells
                for c, ctext in enumerate(row):
                    cells[c].paragraphs[0].text = ''; add_inline(cells[c].paragraphs[0], ctext)
            doc.add_paragraph()
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------- QRC emitter
def emit_qrc(qrc_path, html_name, img_subdir):
    """Embed the HTML + every screenshot in docs/manual/<img_subdir> under :/manual/."""
    rel = os.path.relpath(MANUAL, os.path.dirname(qrc_path)).replace('\\', '/')
    lines = ['<RCC>', '    <qresource prefix="/manual">',
             '        <file alias="%s">%s/%s</file>' % (html_name, rel, html_name)]
    img_dir = os.path.join(MANUAL, img_subdir)
    for png in sorted(os.listdir(img_dir)) if os.path.isdir(img_dir) else []:
        if png.lower().endswith('.png'):
            lines.append('        <file alias="%s/%s">%s/%s/%s</file>' % (img_subdir, png, rel, img_subdir, png))
    lines += ['    </qresource>', '</RCC>', '']
    with open(qrc_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    return qrc_path


def main():
    for md, html_out, docx_out, title, qrc, img_sub in MANUALS:
        blocks = parse_markdown(os.path.join(MANUAL, md))
        h = emit_html(blocks, os.path.join(MANUAL, html_out), title)
        print('wrote', h, os.path.getsize(h), 'bytes')
        try:
            d = emit_docx(blocks, os.path.join(MANUAL, docx_out), MANUAL, title)
            print('wrote', d, os.path.getsize(d), 'bytes')
        except ImportError:
            print('python-docx not installed; skipped', docx_out)
        q = emit_qrc(os.path.normpath(qrc), html_out, img_sub)
        print('wrote', q)


if __name__ == '__main__':
    main()
