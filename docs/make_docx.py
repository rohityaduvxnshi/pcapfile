#!/usr/bin/env python3
"""Convert docs/USER_MANUAL.md into docs/USER_MANUAL.docx with embedded screenshots.
Uses python-docx only (offline). Handles the Markdown constructs the manual uses:
headings, paragraphs (with **bold**/`code`/links), pipe tables, images, blockquotes,
bullet lists, numbered lists, fenced code blocks (ASCII diagrams) and horizontal rules."""
import os, re, struct
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
MD   = os.path.join(HERE, 'USER_MANUAL.md')
OUT  = os.path.join(HERE, 'USER_MANUAL.docx')

def png_size(path):
    try:
        with open(path, 'rb') as f:
            head = f.read(24)
        if head[:8] == b'\x89PNG\r\n\x1a\n':
            w, h = struct.unpack('>II', head[16:24])
            return w, h
    except Exception:
        pass
    return None

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# ---- inline formatting (**bold**, `code`, [text](url)) -------------------------
INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')

def add_inline(par, text):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`'):
            r = par.add_run(tok[1:-1]); r.font.name = 'Consolas'
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            mm = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok)
            label, url = mm.group(1), mm.group(2)
            par.add_run(label + (' (' + url + ')' if url.startswith('http') else ''))
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])

def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'BBBBBB')
    pbdr.append(bottom); pPr.append(pbdr)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def split_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]

def is_sep(line):
    cells = split_row(line)
    return len(cells) > 0 and all(re.fullmatch(r':?-{2,}:?', c) for c in cells)

# ---- read & parse --------------------------------------------------------------
with open(MD, 'r', encoding='utf-8') as f:
    lines = [l.rstrip('\n') for l in f]

i, n = 0, len(lines)
para_buf = []

def flush_para():
    global para_buf
    if para_buf:
        text = ' '.join(s.strip() for s in para_buf).strip()
        if text:
            add_inline(doc.add_paragraph(), text)
        para_buf = []

IMG = re.compile(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$')

while i < n:
    line = lines[i]
    stripped = line.strip()

    # fenced code block
    if stripped.startswith('```'):
        flush_para()
        i += 1
        code = []
        while i < n and not lines[i].strip().startswith('```'):
            code.append(lines[i]); i += 1
        i += 1  # skip closing fence
        for cl in code:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(cl if cl else ' ')
            r.font.name = 'Consolas'; r.font.size = Pt(8.5)
        doc.add_paragraph()  # gap after diagram
        continue

    # image on its own line
    mimg = IMG.match(line)
    if mimg:
        flush_para()
        alt, path = mimg.group(1), mimg.group(2)
        fp = os.path.normpath(os.path.join(HERE, path))
        if os.path.exists(fp):
            sz = png_size(fp)
            width = None
            if sz:
                inch = sz[0] / 96.0
                width = Inches(min(inch, 6.3))
            try:
                doc.add_picture(fp, width=width) if width else doc.add_picture(fp)
            except Exception as e:
                doc.add_paragraph('[image: %s]' % path)
            cap = doc.add_paragraph()
            cr = cap.add_run(alt); cr.italic = True; cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        else:
            doc.add_paragraph('[missing image: %s]' % path)
        i += 1
        continue

    # heading
    mh = re.match(r'^(#{1,6})\s+(.*)$', line)
    if mh:
        flush_para()
        level = len(mh.group(1))
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', mh.group(2))
        text = re.sub(r'`([^`]+)`', r'\1', text)
        doc.add_heading(text, level=0 if level == 1 else min(level - 1, 9))
        i += 1
        continue

    # horizontal rule
    if re.fullmatch(r'(-{3,}|\*{3,}|_{3,})', stripped):
        flush_para(); hrule(); i += 1; continue

    # table
    if stripped.startswith('|') and i + 1 < n and is_sep(lines[i + 1]):
        flush_para()
        header = split_row(line)
        ncol = len(header)
        body = []
        i += 2  # skip header + separator
        while i < n and lines[i].strip().startswith('|'):
            cells = split_row(lines[i])
            if len(cells) < ncol: cells += [''] * (ncol - len(cells))
            body.append(cells[:ncol]); i += 1
        tbl = doc.add_table(rows=1, cols=ncol)
        tbl.style = 'Table Grid'
        tbl.autofit = True
        for c, htext in enumerate(header):
            cell = tbl.rows[0].cells[c]
            shade(cell, 'DDEBFF')
            cell.paragraphs[0].text = ''
            add_inline(cell.paragraphs[0], htext)
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for row in body:
            cells = tbl.add_row().cells
            for c, ctext in enumerate(row):
                cells[c].paragraphs[0].text = ''
                add_inline(cells[c].paragraphs[0], ctext)
        doc.add_paragraph()
        continue

    # blockquote (join consecutive > lines)
    if stripped.startswith('>'):
        flush_para()
        buf = []
        while i < n and lines[i].strip().startswith('>'):
            buf.append(re.sub(r'^\s*>\s?', '', lines[i])); i += 1
        text = ' '.join(s.strip() for s in buf).strip()
        p = doc.add_paragraph(style='Quote')
        add_inline(p, text)
        continue

    # bullet list item (+ indented continuation lines)
    if re.match(r'^[-*]\s+', line):
        flush_para()
        content = re.sub(r'^[-*]\s+', '', line)
        i += 1
        while i < n and re.match(r'^\s{2,}\S', lines[i]) and not re.match(r'^\s*[-*]\s+', lines[i]):
            content += ' ' + lines[i].strip(); i += 1
        p = doc.add_paragraph(style='List Bullet')
        add_inline(p, content)
        continue

    # numbered list item -> keep original number as plain paragraph
    mnum = re.match(r'^(\d+)\.\s+(.*)$', line)
    if mnum:
        flush_para()
        content = mnum.group(2)
        i += 1
        while i < n and re.match(r'^\s{3,}\S', lines[i]) and not re.match(r'^\s*\d+\.\s+', lines[i]):
            content += ' ' + lines[i].strip(); i += 1
        p = doc.add_paragraph()
        rr = p.add_run(mnum.group(1) + '. ')
        add_inline(p, content)
        continue

    # blank line -> paragraph break
    if stripped == '':
        flush_para(); i += 1; continue

    # plain text line -> accumulate
    para_buf.append(line); i += 1

flush_para()
doc.save(OUT)
print('wrote', OUT, os.path.getsize(OUT), 'bytes')
